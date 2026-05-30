import os
from docx import Document
from fpdf import FPDF


def create_resume_docx(
    user_details,
    summary,
    experience,
    projects,
    skills,
    certificates
):
    """
    Generates resume.docx and resume.pdf
    """

    # Get current file directory
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    docx_path = os.path.join(BASE_DIR, "resume.docx")
    pdf_path = os.path.join(BASE_DIR, "resume.pdf")

    # ==========================
    # DOCX GENERATION
    # ==========================

    doc = Document()

    # User Details
    doc.add_heading(user_details["name"], 0)
    doc.add_paragraph(
        f"{user_details['email']} | "
        f"{user_details['phone_number']} | "
        f"{user_details['address']}"
    )

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    # Experience
    doc.add_heading("Experience", level=1)

    for exp in experience:
        doc.add_heading(exp["company"], level=2)
        doc.add_paragraph(
            f"{exp['role']} | {exp['dates']}"
        )
        doc.add_paragraph(exp["achievements"])

    # Projects
    doc.add_heading("Projects", level=1)

    for project in projects:
        doc.add_heading(project["name"], level=2)
        doc.add_paragraph(
            f"{project['description']} | "
            f"{project['technologies']}"
        )

    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(skills))

    # Certificates
    doc.add_heading("Certificates", level=1)

    for cert in certificates:
        doc.add_heading(cert["name"], level=2)
        doc.add_paragraph(cert["description"])

    # Save DOCX
    doc.save(docx_path)

    # ==========================
    # PDF GENERATION
    # ==========================

    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Arial", size=12)

    # User Details
    pdf.cell(0, 10, user_details["name"], ln=True, align="C")

    pdf.multi_cell(
        0,
        10,
        f"{user_details['email']} | "
        f"{user_details['phone_number']} | "
        f"{user_details['address']}"
    )

    pdf.ln(3)

    # Summary
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 10, "Summary", ln=True)

    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, summary)

    # Experience
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 10, "Experience", ln=True)

    pdf.set_font("Arial", size=12)

    for exp in experience:
        pdf.cell(0, 10, exp["company"], ln=True)
        pdf.cell(
            0,
            10,
            f"{exp['role']} | {exp['dates']}",
            ln=True
        )
        pdf.multi_cell(0, 10, exp["achievements"])

    # Projects
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 10, "Projects", ln=True)

    pdf.set_font("Arial", size=12)

    for project in projects:
        pdf.cell(0, 10, project["name"], ln=True)
        pdf.multi_cell(
            0,
            10,
            f"{project['description']} | "
            f"{project['technologies']}"
        )

    # Skills
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 10, "Skills", ln=True)

    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, ", ".join(skills))

    # Certificates
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 10, "Certificates", ln=True)

    pdf.set_font("Arial", size=12)

    for cert in certificates:
        pdf.cell(0, 10, cert["name"], ln=True)
        pdf.multi_cell(0, 10, cert["description"])

    # Save PDF
    pdf.output(pdf_path)

    print(f"DOCX saved at: {docx_path}")
    print(f"PDF saved at : {pdf_path}")


# ===================================
# TESTING
# ===================================

if __name__ == "__main__":

    user_details = {
        "name": "Darshan Salunke",
        "email": "darshan@gmail.com",
        "phone_number": "8208313549",
        "address": "Pune, Maharashtra"
    }

    summary = (
        "Backend Developer skilled in Python, Django, "
        "FastAPI, PostgreSQL and Generative AI."
    )

    experience = [
        {
            "company": "Bynry Inc",
            "role": "Backend Developer Intern",
            "dates": "Aug 2025 - Present",
            "achievements": (
                "Built REST APIs, optimized database queries, "
                "and worked on scalable backend systems."
            )
        }
    ]

    projects = [
        {
            "name": "PDF Q&A System",
            "description": (
                "Real-time PDF question answering application"
            ),
            "technologies": (
                "FastAPI, Transformers, PyTorch, SQLite"
            )
        }
    ]

    skills = [
        "Python",
        "Django",
        "FastAPI",
        "PostgreSQL",
        "Generative AI"
    ]

    certificates = [
        {
            "name": "AWS Cloud Practitioner",
            "description": "Cloud computing fundamentals"
        }
    ]

    create_resume_docx(
        user_details,
        summary,
        experience,
        projects,
        skills,
        certificates
    )

    print("Resume generated successfully!")